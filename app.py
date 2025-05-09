# Updated Streamlit app to reflect new content structure and clean errors
import streamlit as st
import pandas as pd
import openai
from dotenv import load_dotenv
import os
import json

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

st.set_page_config(page_title="Product Listing Automation", layout="wide")
st.title("🧠 Product Listing Automation Tool")

st.subheader("📥 Upload KLD Sheet")
uploaded_file = st.file_uploader("Upload your KLD Excel file", type=["xlsx"])

if uploaded_file:
    df = pd.read_excel(uploaded_file)
    df.columns = df.columns.str.strip()
    st.success("KLD sheet loaded successfully!")

    if 'Product Name' in df.columns:
        product_info = df.iloc[0]

        st.subheader("🔹 Product Information")
        product_table = {
            "Brand Name": product_info['Brand Name'],
            "Product Name": product_info['Product Name'],
            "USPs Front": product_info['USPs Front'],
            "Ingredients": product_info['Ingredients'],
            "Claims": product_info['Claims'],
            "How to use it?": product_info['How to use it?'],
            "Appropriate Age": product_info['Appropriate Age'],
            "Net Weight": product_info['Net Weight'],
            "MRP": f"₹{product_info['MRP']}",
            "Pack Contains": product_info['Particulars']
        }
        st.table(pd.DataFrame(product_table.items(), columns=["Field", "Value"]))

        def generate_section(section_name, instruction):
            try:
                completion = openai.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "You are a professional ecommerce copywriter."},
                        {"role": "user", "content": instruction}
                    ],
                    temperature=0.7
                )
                return completion.choices[0].message.content.strip()
            except Exception as e:
                st.error(f"Error generating {section_name}: {e}")
                return ""

        def text_box(label, key, height):
            if key in st.session_state:
                st.text_area(label, st.session_state[key], height=height)

        # Amazon Content Generation
        st.session_state.setdefault('title', '')
        if st.button("Generate Product Title"):
            title_prompt = f"""
Write an Amazon product title (max 200 characters) for a high-converting listing.
Use this format: Brand + Product Type + Keywords + Claims.
Do not include size, weight, or volume in the title.

Product: {product_info['Product Name']}
Brand: {product_info['Brand Name']}
USPs: {product_info['USPs Front']}
Ingredients: {product_info['Ingredients']}
Claims: {product_info['Claims']}
"""
            st.session_state['title'] = generate_section("title", title_prompt)
        text_box("Amazon Product Title", 'title', 100)
        st.subheader("🛒 Amazon Content")
        st.session_state.setdefault('bullets', '')
        if st.button("Generate Bullet Points"):
            bullet_prompt = f"""
Write 7 optimized Amazon bullet points. Each bullet should follow this format:
BENEFIT IN CAPS: Followed by a clear, compelling benefit (250–300 characters).

Avoid mentioning price, value for money, discounts, or any numerical cost-related details. Focus only on features, results, usage, or ingredient benefits.

Product: {product_info['Product Name']}
Brand: {product_info['Brand Name']}
USPs: {product_info['USPs Front']}
Ingredients: {product_info['Ingredients']}
Claims: {product_info['Claims']}
How to Use: {product_info['How to use it?']}
Pack Contains: {product_info['Particulars']}
MRP: ₹{product_info['MRP']}
            """
            st.session_state['bullets'] = generate_section("bullet points", bullet_prompt)
        text_box("Generated Bullet Points", 'bullets', 200)

        st.session_state.setdefault('description', '')
        if st.button("Generate Description"):
            desc_prompt = f"""
Write an Amazon HTML product description (max 400 words).
Use 2 paragraphs, <p> and <br> tags for light formatting.

Product: {product_info['Product Name']}
Brand: {product_info['Brand Name']}
USPs: {product_info['USPs Front']}
Ingredients: {product_info['Ingredients']}
Claims: {product_info['Claims']}
How to Use: {product_info['How to use it?']}
            """
            st.session_state['description'] = generate_section("description", desc_prompt)
        text_box("HTML Description", 'description', 150)

        st.session_state.setdefault('shopify', '')
        if st.button("Generate Shopify Description"):
            shopify_prompt = f"""
Write a Shopify product description using a friendly, informative tone.
Include small paragraphs, bullet points, and headings. Length: 1500–2000 characters.

Product: {product_info['Product Name']}
Brand: {product_info['Brand Name']}
USPs: {product_info['USPs Front']}
Ingredients: {product_info['Ingredients']}
Claims: {product_info['Claims']}
How to Use: {product_info['How to use it?']}
Pack Contains: {product_info['Particulars']}
MRP: ₹{product_info['MRP']}
            """
            st.session_state['shopify'] = generate_section("Shopify description", shopify_prompt)
        text_box("Shopify Description", 'shopify', 200)

        st.session_state.setdefault('hero', '')
        if st.button("Generate Hero Prompts"):
            if product_info['Brand Name'] == "Urbanyog":
                hero_prompt = f"""
                    Generate a series of 7 hero image prompts for a beauty product, each designed to be a standalone visual in the following sequence.  
                    Each image must be exactly 1500 x 1500 pixels. For each image, include specific visual/creative directions as outlined.

                    1. Hero Image: What does it do / Differentiation / Before-After  
                    - Show the product with a clear before-and-after comparison, highlighting the main benefit or transformation.  
                    - Use high-resolution, professional imagery with a clean, softly colored, on-brand background.  
                    - Overlay concise benefit-driven text or icons.  
                    - Ensure the product is the main focal point.

                    2. USP  
                    - Focus on the unique selling proposition that sets this product apart.  
                    - Use bold text or a badge to highlight the USP.  
                    - Incorporate lifestyle or in-use imagery for emotional connection.

                    3. Ingredients & Their Use  
                    - Visually showcase key ingredients with small icons or illustrations.  
                    - Briefly mention each ingredient’s benefit.  
                    - Use a soft, inviting color palette.

                    4. Science Behind USP  
                    - Include science-related visuals (molecules, lab glassware, dermatologist icons) to reinforce efficacy.  
                    - Overlay a short, clear explanation of the science behind the USP.

                    5. Comparison / Do's and Don'ts  
                    - Create a side-by-side comparison with alternatives or a clear Do’s and Don’ts visual.  
                    - Use simple icons and minimal text for clarity.

                    6. How to Use  
                    - Present step-by-step usage instructions in 3–4 crisp, numbered icons or steps.  
                    - Keep instructions short, clear, and visually engaging.

                    7. Safe for All Hair/Skin Types  
                    - Use icons or imagery showing diversity in models (different skin/hair types).  
                    - Add a clear statement or badge indicating suitability for all.

                    General Visual/Creative Directions for All Images:  
                    - Each image must be exactly 1500 x 1500 pixels.  
                    - Use high-resolution, professional images.  
                    - Clean, uncluttered, on-brand backgrounds.  
                    - Overlay concise, benefit-driven text or icons.  
                    - Maintain a consistent style, color palette, and font across all images.  
                    - Optimize for both desktop and mobile screens.  
                    - Avoid clutter; keep visuals focused and easy to scan.  
                    - Do not include any call-to-action (CTA) buttons.  
                    - Include relevant props that complement the product and provide context, without distracting from the main subject.  
                    - Feature models where appropriate, ensuring diversity in skin and hair types to show inclusivity and real-world results.  
                    - Use professional lighting setups (such as three-point lighting, softboxes, umbrellas, or ring lights) for soft, even illumination and to minimize harsh shadows.  
                    - Capture the product from optimal angles, including close-ups for detail and lifestyle shots for context.  
                    - Ensure the product is always the main focal point, with props and models used to enhance the story.

                    Product details:  
                    Product: {product_info['Product Name']}  
                    USP: {product_info['USPs Front']}  
                    Ingredients: {product_info['Ingredients']}  
                    How to use: {product_info['How to use it?']}  
                    Target: {product_info.get('Target Audience', '')}  
                    Other details: {product_info.get('Other details', '')}
                    """
            elif product_info['Brand Name'] == "MakeMeeBold":
                hero_prompt = f"""
                Generate a series of 7 hero image prompts for an electronics product, each designed to be a standalone visual in the following sequence.  
                Each image must be exactly 1500 x 1500 pixels. For each image, include specific visual/creative directions as outlined.

                1. Hero Image: What does it do / Differentiation / Before-After  
                - Show the product in action or with a before-and-after or comparison visual.  
                - Use high-resolution, professional imagery with a clean, on-brand background.  
                - Overlay concise benefit-driven text or icons.

                2. Main USP  
                - Highlight the primary unique selling point with bold text or a badge.  
                - Use dynamic angles and close-ups to draw attention.

                3. Other USPs  
                - Present additional USPs or features with icons or short text.  
                - Arrange features for easy scanning.

                4. Comparison  
                - Create a side-by-side comparison with other products or brands.  
                - Use simple graphics and minimal text.

                5. How to Use  
                - Show 2–4 easy-to-follow steps or icons for product usage.  
                - Keep instructions short and visually engaging.

                6. All Hair Type (if relevant)  
                - Use a badge or icon to indicate compatibility with all hair types (for grooming devices).  
                - Show diverse models if applicable.

                7. What's in the Box / Customer Care  
                - Display all included accessories in a neat flat-lay or organized arrangement.  
                - Add customer care contact or warranty info in a clear, non-intrusive way.

                General Visual/Creative Directions for All Images:  
                - Each image must be exactly 1500 x 1500 pixels.  
                - Use high-resolution, professional images.  
                - Clean, uncluttered, on-brand backgrounds.  
                - Overlay concise, benefit-driven text or icons.  
                - Consistent style, color palette, and font across all images.  
                - Optimize for both desktop and mobile screens.  
                - Avoid clutter; keep visuals focused and easy to scan.  
                - Do not include any call-to-action (CTA) buttons.  
                - Include relevant props that complement the product and provide context, without distracting from the main subject.  
                - Feature models where appropriate (e.g., for scale or lifestyle context).  
                - Use professional lighting setups (such as three-point lighting, softboxes, umbrellas, or ring lights) for soft, even illumination and to minimize harsh shadows.  
                - Capture the product from optimal angles, including close-ups for detail and lifestyle shots for context.  
                - Ensure the product is always the main focal point, with props and models used to enhance the story.

                Product details:  
                Product: {product_info['Product Name']}  
                Main USP: {product_info['USPs Front']}  
                Other USPs: {product_info.get('Other USPs', '')}  
                How to use: {product_info['How to use it?']}  
                Box contents: {product_info.get('Particulars', '')}  
                Customer care: {product_info.get('Customer Care', '')}  
                Other details: {product_info.get('Other details', '')}
                """
            else:
                hero_prompt = "Category not supported. Please specify 'beauty' or 'electronics' in the product info."

            st.session_state['hero'] = generate_section("hero image prompts", hero_prompt)
        text_box("Hero Image Prompts", 'hero', 180)

        st.session_state.setdefault('a_plus', '')
        if st.button("Generate A+ Prompts"):
            a_plus_prompt = f"""
Generate 7 Amazon A+ content image prompts in this format:

Image 1:
• Visual: 1464x600 (desktop) / 600x450 (mobile). Clean layout, brand, benefits.
• Text: Marketing headline
• Supporting Text: Short supporting copy

Product: {product_info['Product Name']}
USPs: {product_info['USPs Front']}
How to Use: {product_info['How to use it?']}
Claims: {product_info['Claims']}
Ingredients: {product_info['Ingredients']}
            """
            st.session_state['a_plus'] = generate_section("A+ image prompts", a_plus_prompt)
        text_box("A+ Image Prompts", 'a_plus', 180)

        # Website Content Generation
        if st.button("🧠 Generate Full Website Content"):
            full_web_prompt = f"""
Generate the following website content in a structured format:

1. 7 Bullet Points – focus on customer benefits, pain points, or product uniqueness.
2. Description – 2 paragraphs, warm and benefit-driven tone, max 400 words.
3. USP Points – concise value-driven phrases (2–4 words each).
4. What do you get – brief explanation of what comes in the package.
5. How to use – easy-to-follow, friendly, step-by-step instructions.
6. 6 FAQs – with short, helpful answers.
7. 15 Customer Reviews – first 2 as story-style testimonials, remaining 13 as short user opinions.

Product Name: {product_info['Product Name']}
Brand: {product_info['Brand Name']}
USPs / Features: {product_info['USPs Front']}
Ingredients: {product_info['Ingredients']}
Claims: {product_info['Claims']}
How to Use: {product_info['How to use it?']}
What’s Included: {product_info['Particulars']}
MRP: ₹{product_info['MRP']}
            """
            web_full = generate_section("Website Content", full_web_prompt)
            st.session_state['web_full'] = web_full
        text_box("Full Website Content", 'web_full', 500)
        st.subheader("🌐 Website Content")

        st.session_state.setdefault('web_bullets', '')
        if st.button("Generate Website Bullet Points"):
            web_bullet_prompt = f"""
Write 7 bullet points for website use. Each point should focus on product benefits, customer problems, or emotional triggers.

Product: {product_info['Product Name']}
USPs: {product_info['USPs Front']}
"""
            st.session_state['web_bullets'] = generate_section("website bullets", web_bullet_prompt)
        text_box("Website Bullet Points", 'web_bullets', 200)

        st.session_state.setdefault('web_description', '')
        if st.button("Generate Website Description"):
            web_desc_prompt = f"""
Write a product description for the website (max 400 words). Use a warm, benefit-oriented tone and structure it with two paragraphs.

Product: {product_info['Product Name']}
USPs: {product_info['USPs Front']}
Claims: {product_info['Claims']}
Ingredients: {product_info['Ingredients']}
"""
            st.session_state['web_description'] = generate_section("website description", web_desc_prompt)
        text_box("Website Description", 'web_description', 200)

        st.session_state.setdefault('usp', '')
        if st.button("Generate USP"):
            usp_prompt = f"""
Generate exactly 6 concise USPs for a product listing.
Guidelines:
- Each USP must be a short, impactful phrase of 2–4 words.
- Do not include full sentences.
- Do not include explanations.
- Format strictly as a bullet list like:
  - Frizz-control Formula
  - Lightweight & Non-Greasy
  - Safe for Daily Use

Product: {product_info['Product Name']}
USPs: {product_info['USPs Front']}
"""
            st.session_state['usp'] = generate_section("USP", usp_prompt)
        text_box("USP", 'usp', 100)

        st.session_state.setdefault('what_you_get', '')
        if st.button("Generate 'What Do You Get'"):
            get_prompt = f"""
Describe what the customer will receive when they purchase the product. Mention packaging, quantity, and any bonuses.

Product: {product_info['Product Name']}
"""
            st.session_state['what_you_get'] = generate_section("What Do You Get", get_prompt)
        text_box("What Do You Get", 'what_you_get', 100)

        st.session_state.setdefault('how_to_use', '')
        if st.button("Generate 'How to Use'"):
            how_to_prompt = f"""
Describe how to use this product step-by-step in a friendly, instructional tone.

Product: {product_info['Product Name']}
Instructions: {product_info['How to use it?']}
"""
            st.session_state['how_to_use'] = generate_section("How to Use", how_to_prompt)
        text_box("How to Use This", 'how_to_use', 100)

        st.session_state.setdefault('faqs', '')
        if st.button("Generate FAQs"):
            faqs_prompt = f"""
Write 6 frequently asked questions and concise answers for this Amazon product listing.
Guidelines:
- Do NOT include questions about certifications, result timelines, or return/refund policies.
- Avoid questions like "What is it?" or "How do I use it?" since those are covered elsewhere.
- Focus on practical use, compatibility, safety (general, not certifications), storage, frequency, product care, and common customer concerns.
- Keep answers under 150 characters.
- Use clear, customer-friendly language.
- Naturally include relevant keywords where possible, but do not keyword stuff.
- Do not mention price, promotions, or other brands.

Product: {product_info['Product Name']}
Key Features: {product_info['USPs Front']}
Ingredients: {product_info['Ingredients']}
How to Use: {product_info['How to use it?']}

"""
            st.session_state['faqs'] = generate_section("FAQs", faqs_prompt)
        text_box("FAQs", 'faqs', 200)

        st.session_state.setdefault('reviews', '')
        if st.button("Generate Reviews"):
            review_prompt = f"""
Write 15 customer reviews for this product.
- The first 2 should be long story-style testimonials.
- The remaining 13 should be short, specific, and highlight different use cases or outcomes.
- All reviewer names should sound authentic and be common Indian names.

Product: {product_info['Product Name']}
"""
            st.session_state['reviews'] = generate_section("Reviews", review_prompt)
        text_box("Customer Reviews", 'reviews', 300)

        # Export or clear
        if st.button("📥 Download Output as Word"):
            import docx
            from docx.shared import Pt
            from io import BytesIO

            doc = docx.Document()
            doc.add_heading(f"Product Listing Content – {product_info['Product Name']}", 0)

            sections = [
                ('Amazon Bullet Points', 'bullets'),
                ('Amazon Description', 'description'),
                ('Shopify Description', 'shopify'),
                ('Hero Image Prompts', 'hero'),
                ('A+ Image Prompts', 'a_plus'),
                ('Website Bullet Points', 'web_bullets'),
                ('Website Description', 'web_description'),
                ('USP', 'usp'),
                ('What Do You Get', 'what_you_get'),
                ('How to Use', 'how_to_use'),
                ('FAQs', 'faqs'),
                ('Customer Reviews', 'reviews')
            ]

            for title, key in sections:
                content = st.session_state.get(key, '')
                if content:
                    doc.add_heading(title, level=1)
                    doc.add_paragraph(content, style='Normal')

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button(
                label="Download Word Document",
                data=buffer,
                file_name="product_listing.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        if st.button("🧹 Clear All"):
            for key in ['bullets', 'description', 'shopify', 'hero', 'a_plus', 'web_bullets', 'web_description', 'usp', 'what_you_get', 'how_to_use', 'faqs', 'reviews', 'web_full']:
                st.session_state[key] = ''

        if st.button("📥 Download Output as JSON"):
            output = {k: st.session_state.get(k, '') for k in ['bullets', 'description', 'shopify', 'hero', 'a_plus', 'web_bullets', 'web_description', 'usp', 'what_you_get', 'how_to_use', 'faqs', 'reviews', 'web_full']}
            json_str = json.dumps(output, indent=2)
            st.download_button("Download JSON", data=json_str, file_name="listing_content.json", mime="application/json")

    else:
        st.error("'Product Name' column not found. Please check your sheet headers.")
else:
    st.info("Please upload your KLD sheet to get started.")
